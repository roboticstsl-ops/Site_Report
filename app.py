import os
import copy
import base64
import io
import subprocess
import tempfile

import fitz  # PyMuPDF

from flask import Flask, render_template, request, send_file
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches
from docx.text.paragraph import Paragraph as DocxParagraph

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024  # 50 MB

TEMPLATE_PATH = os.path.join(os.path.dirname(__file__), 'SiteReportTemplate (1).docx')

CHECKLIST_ITEMS = [
    "Choose elevator(s) to be used (service/resident)",
    "Test connectivity at elevator (Wi-Fi / LTE)",
    "Confirm elevator compatibility (door/space)",
    "Identify docking location(s)",
    "Confirm power availability at docking",
    "Confirm access doors or obstacles along robot route",
]

CHECKLIST_KEYS = [
    'ElevatorSelected',
    'ConnectivityOK',
    'ElevatorCompatibility',
    'DockingIdentified',
    'PowerAvailable',
    'AccessOK',
]

OWNERS = ["Building", "Client", "TII", "Vendor", "TBD"]


# ── Placeholder replacement helpers ──────────────────────────────────────────

def _replace_in_element(element, replacements):
    """
    Replace <<Key>> markers in all w:p elements within `element`.
    Merges split runs per paragraph before replacing, so even keys
    that Word has split across multiple w:r elements are handled.
    """
    for para in element.iter(qn('w:p')):
        runs = para.findall('.//' + qn('w:r'))
        if not runs:
            continue

        # Reconstruct full paragraph text from all runs
        texts = []
        for r in runs:
            t = r.find(qn('w:t'))
            texts.append(t.text or '' if t is not None else '')
        full_text = ''.join(texts)

        # Skip paragraphs that contain no placeholders we care about
        if not any(f'<<{k}>>' in full_text for k in replacements):
            continue

        # Merge all runs into the first run to fix any split placeholder
        first_run = runs[0]
        first_t = first_run.find(qn('w:t'))
        if first_t is None:
            continue
        first_t.text = full_text
        if full_text.startswith(' ') or full_text.endswith(' '):
            first_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')

        for run in runs[1:]:
            parent = run.getparent()
            if parent is not None:
                parent.remove(run)

        # Do text replacement on the merged run
        for key, value in replacements.items():
            first_t.text = (first_t.text or '').replace(f'<<{key}>>', value)


def _replace_caption_placeholders(element, captions):
    """
    Replace <<Image1Caption>> markers one by one (in document order)
    so each of the 4 identical placeholder names gets a different caption.
    """
    cap_idx = 0
    for para in element.iter(qn('w:p')):
        if cap_idx >= len(captions):
            break
        runs = para.findall('.//' + qn('w:r'))
        texts = [
            (r.find(qn('w:t')).text or '')
            for r in runs
            if r.find(qn('w:t')) is not None
        ]
        full_text = ''.join(texts)
        if '<<Image1Caption>>' not in full_text:
            continue

        # Merge runs and replace just once
        first_run = runs[0]
        first_t = first_run.find(qn('w:t'))
        if first_t is None:
            continue
        first_t.text = full_text
        for run in runs[1:]:
            p = run.getparent()
            if p is not None:
                p.remove(run)

        first_t.text = (first_t.text or '').replace(
            '<<Image1Caption>>', captions[cap_idx], 1
        )
        cap_idx += 1


def _insert_images(doc, element, images):
    """
    Find every paragraph containing '[ Image Placeholder ]' inside `element`
    and replace it with the corresponding uploaded image.
    Returns a list of original image bytes (or None) in slot order,
    so the caller can post-process the PDF with full-quality originals.
    """
    orig_bytes_list = []
    slot_idx = 0
    for para_el in element.iter(qn('w:p')):
        if slot_idx >= len(images):
            break
        runs = para_el.findall('.//' + qn('w:r'))
        text = ''.join(
            (r.find(qn('w:t')).text or '')
            for r in runs
            if r.find(qn('w:t')) is not None
        )
        if '[ Image Placeholder ]' not in text:
            continue

        img_data = images[slot_idx]
        slot_idx += 1

        if not img_data.get('data_uri'):
            orig_bytes_list.append(None)
            # Clear placeholder text but leave cell empty
            for r in list(para_el.findall('.//' + qn('w:r'))):
                r.getparent().remove(r)
            continue

        # Decode image bytes
        _, b64 = img_data['data_uri'].split(',', 1)
        img_bytes = base64.b64decode(b64)
        orig_bytes_list.append(img_bytes)  # keep original for PDF post-processing

        # Clear all existing runs from the paragraph
        for r in list(para_el.findall('.//' + qn('w:r'))):
            r.getparent().remove(r)
        # Also remove any SDT wrappers
        for sdt in list(para_el.findall('.//' + qn('w:sdt'))):
            sdt.getparent().remove(sdt)

        # Use python-docx Paragraph wrapper to add a picture run
        para_obj = DocxParagraph(para_el, doc)
        run = para_obj.add_run()
        run.add_picture(io.BytesIO(img_bytes), width=Inches(2.60), height=Inches(2.60))

    return orig_bytes_list


def _set_cell_text(cell_el, text):
    """Replace all content in a table cell with a single plain-text paragraph."""
    for p in list(cell_el.findall(qn('w:p'))):
        cell_el.remove(p)
    p = OxmlElement('w:p')
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = text or ''
    if t.text.startswith(' ') or t.text.endswith(' '):
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    r.append(t)
    p.append(r)
    cell_el.append(p)


def _add_extra_change_rows(assess_el, changes):
    """
    For changes beyond the 2 slots in the template, clone the last data row
    and fill it with plain text so unlimited changes are supported.
    """
    if len(changes) <= 2:
        return

    # Find the site-changes nested table (has "Change" + "Owner" headers)
    # Skip assess_el itself (iter() includes the root element)
    changes_tbl = None
    for tbl in assess_el.iter(qn('w:tbl')):
        if tbl is assess_el:
            continue
        rows = tbl.findall(qn('w:tr'))
        if len(rows) >= 2:
            hdr_text = ''.join(t.text or '' for t in rows[0].iter(qn('w:t')))
            if 'Change' in hdr_text and 'Owner' in hdr_text:
                changes_tbl = tbl
                break
    if changes_tbl is None:
        return

    data_rows = changes_tbl.findall(qn('w:tr'))[1:]  # skip header
    if not data_rows:
        return

    template_row = data_rows[-1]  # clone from last existing data row

    for change in changes[2:]:
        new_row = copy.deepcopy(template_row)
        # Clear any remaining placeholders in the cloned row
        for t_el in new_row.iter(qn('w:t')):
            if '<<' in (t_el.text or ''):
                t_el.text = ''

        cells = new_row.findall(qn('w:tc'))
        if len(cells) >= 1:
            _set_cell_text(cells[0], change.get('change', ''))
        if len(cells) >= 3:
            _set_cell_text(cells[2], change.get('notes', ''))

        template_row.addnext(new_row)
        template_row = new_row  # next clone appends after this


# ── Document building ─────────────────────────────────────────────────────────

def _add_page_break_after(element):
    """Insert a page-break paragraph immediately after `element`."""
    p = OxmlElement('w:p')
    r = OxmlElement('w:r')
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    r.append(br)
    p.append(r)
    element.addnext(p)
    return p


def build_document(header, towers):
    """
    Load the DOCX template, fill header once, then fill/duplicate the
    tower section for each tower.  Returns a filled Document object.
    """
    doc = Document(TEMPLATE_PATH)
    body = doc.element.body
    children = list(body)

    # Body order: [0]=header table, [1]=empty para,
    #             [2]=tower heading para, [3]=assessment table,
    #             [4]=empty para, [5]=sectPr
    header_tbl   = children[0]
    tower_hdg_tpl = children[2]   # "Tower 1: <<Tower>>"
    assess_tbl_tpl = children[3]  # big assessment table

    # Snapshot originals BEFORE any replacement (used as templates for towers 2+)
    tower_hdg_orig   = copy.deepcopy(tower_hdg_tpl)
    assess_tbl_orig  = copy.deepcopy(assess_tbl_tpl)

    # ── Fill header table ──────────────────────────────────────────────────
    _replace_in_element(header_tbl, {
        'Area':            header['area'],
        'Date':            header['date'],
        'Assessor':        header['assessor'],
        'AssessorContact': header['assessor_contact'],
    })

    # ── Fill each tower ────────────────────────────────────────────────────
    last_el = assess_tbl_tpl  # track insertion point for subsequent towers
    all_orig_images = []      # original image bytes, in PDF image order

    for idx, tower in enumerate(towers):
        cl = tower['checklist']
        sc = tower['site_changes']
        imgs = tower.get('images', [])

        repl = {
            'Tower':                     tower['name'],
            'ElevatorSelected':          cl[0]['status'] if len(cl) > 0 else '',
            'ElevatorSelectedNote':      cl[0]['notes']  if len(cl) > 0 else '',
            'ConnectivityOK':            cl[1]['status'] if len(cl) > 1 else '',
            'ConnectivityOKNote':        cl[1]['notes']  if len(cl) > 1 else '',
            'ElevatorCompatibility':     cl[2]['status'] if len(cl) > 2 else '',
            'ElevatorCompatibilityNote': cl[2]['notes']  if len(cl) > 2 else '',
            'DockingIdentified':         cl[3]['status'] if len(cl) > 3 else '',
            'DockingIdentifiedNote':     cl[3]['notes']  if len(cl) > 3 else '',
            'PowerAvailable':            cl[4]['status'] if len(cl) > 4 else '',
            'PowerAvailableNote':        cl[4]['notes']  if len(cl) > 4 else '',
            'AccessOK':                  cl[5]['status'] if len(cl) > 5 else '',
            'AccessOKNote':              cl[5]['notes']  if len(cl) > 5 else '',
            'ElevatorType':              tower['docking']['elevator_type'],
            'DockingNotes':              tower['docking']['notes'],
            'ElevatorDoorWidth':         tower['elevator']['door_width'],
            'ElevatorCabinDepth':        tower['elevator']['cabin_depth'],
            'ElevatorProvider':          tower['elevator']['provider'],
            'ElevatorComments':          tower['elevator']['comments'],
            'Changes1':                  sc[0]['change'] if len(sc) > 0 else '',
            'Changes1Note':              sc[0]['notes']  if len(sc) > 0 else '',
            'Changes2':                  sc[1]['change'] if len(sc) > 1 else '',
            'Changes2Note':              sc[1]['notes']  if len(sc) > 1 else '',
        }
        captions = [img.get('caption', '') for img in imgs[:4]]
        while len(captions) < 4:
            captions.append('')

        if idx == 0:
            # Fill the template's own elements
            _replace_in_element(tower_hdg_tpl, repl)
            _replace_in_element(assess_tbl_tpl, repl)
            _add_extra_change_rows(assess_tbl_tpl, sc)
            _replace_caption_placeholders(assess_tbl_tpl, captions)
            orig = _insert_images(doc, assess_tbl_tpl, imgs)
            last_el = assess_tbl_tpl
        else:
            # Deep-copy from the clean originals (not the already-filled ones)
            pb = _add_page_break_after(last_el)
            new_hdg = copy.deepcopy(tower_hdg_orig)
            pb.addnext(new_hdg)
            new_tbl = copy.deepcopy(assess_tbl_orig)
            new_hdg.addnext(new_tbl)

            _replace_in_element(new_hdg, repl)
            _replace_in_element(new_tbl, repl)
            _add_extra_change_rows(new_tbl, sc)
            _replace_caption_placeholders(new_tbl, captions)
            orig = _insert_images(doc, new_tbl, imgs)
            last_el = new_tbl

        all_orig_images.extend(orig)

    return doc, all_orig_images


def _replace_pdf_images_hq(pdf_bytes: bytes, orig_image_list: list) -> bytes:
    """
    Post-process a PDF produced by LibreOffice to replace its downsampled
    images with the original high-quality versions.

    orig_image_list: list of bytes-or-None, one entry per image slot in
    document order (matching the order images appear in the PDF).
    """
    if not any(b is not None for b in orig_image_list):
        return pdf_bytes  # nothing to do

    doc = fitz.open(stream=pdf_bytes, filetype='pdf')

    # Collect unique image xrefs in document order, noting the first page each appears on
    seen_xrefs = []
    xref_to_page_num = {}
    for page_num in range(len(doc)):
        for img_info in doc[page_num].get_images(full=True):
            xref = img_info[0]
            if xref not in seen_xrefs:
                seen_xrefs.append(xref)
                xref_to_page_num[xref] = page_num

    for i, xref in enumerate(seen_xrefs):
        if i >= len(orig_image_list):
            break
        orig_bytes = orig_image_list[i]
        if orig_bytes is None:
            continue

        # If the image is not JPEG, convert to high-quality JPEG so the PDF
        # stream is always in DCT format (universally supported).
        if orig_bytes[:3] != b'\xff\xd8\xff':
            from PIL import Image as _PIL
            pil = _PIL.open(io.BytesIO(orig_bytes))
            if pil.mode in ('RGBA', 'LA', 'P'):
                pil = pil.convert('RGB')
            buf = io.BytesIO()
            pil.save(buf, format='JPEG', quality=95, subsampling=0)
            orig_bytes = buf.getvalue()

        page = doc[xref_to_page_num[xref]]
        page.replace_image(xref, stream=orig_bytes)

    return doc.tobytes(garbage=4, deflate=True)


def docx_to_pdf(docx_bytes: bytes) -> bytes:
    """Convert a DOCX byte-string to PDF using LibreOffice."""
    with tempfile.TemporaryDirectory() as tmp:
        in_path = os.path.join(tmp, 'report.docx')
        out_path = os.path.join(tmp, 'report.pdf')
        with open(in_path, 'wb') as f:
            f.write(docx_bytes)
        subprocess.run(
            ['libreoffice', '--headless', '--convert-to', 'pdf',
             '--outdir', tmp, in_path],
            check=True, capture_output=True
        )
        with open(out_path, 'rb') as f:
            return f.read()


# ── Form parsing helpers ──────────────────────────────────────────────────────

def parse_tower(form, files, tower_idx):
    t = str(tower_idx)
    tower = {
        'name': form.get(f'tower_{t}_name', f'Tower {tower_idx}'),
        'checklist': [],
        'docking': {
            'elevator_type': form.get(f'tower_{t}_dock_type', 'Service'),
            'notes':         form.get(f'tower_{t}_dock_notes', ''),
        },
        'elevator': {
            'door_width': form.get(f'tower_{t}_elev_width', ''),
            'cabin_depth': form.get(f'tower_{t}_elev_depth', ''),
            'provider':   form.get(f'tower_{t}_elev_provider', ''),
            'comments':   form.get(f'tower_{t}_elev_comments', ''),
        },
        'site_changes': [],
        'images': [],
    }

    for i, _ in enumerate(CHECKLIST_ITEMS):
        tower['checklist'].append({
            'status': form.get(f'tower_{t}_check_{i}_status', 'No'),
            'notes':  form.get(f'tower_{t}_check_{i}_notes', ''),
        })

    change_count = int(form.get(f'tower_{t}_change_count', 0))
    for i in range(change_count):
        change = form.get(f'tower_{t}_change_{i}_text', '')
        notes  = form.get(f'tower_{t}_change_{i}_notes', '')
        if change or notes:
            tower['site_changes'].append({'change': change, 'notes': notes})

    img_count = int(form.get(f'tower_{t}_img_count', 0))
    for j in range(img_count):
        caption = form.get(f'tower_{t}_img_{j}_caption', '')
        file_storage = files.get(f'tower_{t}_img_{j}_file')
        data_uri = None
        if file_storage and file_storage.filename:
            data = file_storage.read()
            if data:
                mime = file_storage.mimetype or 'image/jpeg'
                data_uri = f"data:{mime};base64,{base64.b64encode(data).decode()}"
        tower['images'].append({'caption': caption, 'data_uri': data_uri})

    return tower


# ── Routes ────────────────────────────────────────────────────────────────────

@app.route('/')
def index():
    return render_template('index.html',
                           checklist_items=CHECKLIST_ITEMS,
                           owners=OWNERS)


@app.route('/generate', methods=['POST'])
def generate():
    form  = request.form
    files = request.files

    header = {
        'area':              form.get('area', ''),
        'date':              form.get('date', ''),
        'assessor':          form.get('assessor', ''),
        'assessor_contact':  form.get('assessor_contact', ''),
    }

    active_str = form.get('active_towers', '')
    if active_str:
        active_indices = [int(x) for x in active_str.split(',') if x.strip().isdigit()]
    else:
        active_indices = list(range(1, int(form.get('tower_count', 1)) + 1))
    towers = [parse_tower(form, files, i) for i in active_indices]

    doc, orig_images = build_document(header, towers)

    # Save DOCX to bytes
    buf = io.BytesIO()
    doc.save(buf)
    docx_bytes = buf.getvalue()

    # Convert to PDF via LibreOffice, then restore original image quality
    pdf_bytes = docx_to_pdf(docx_bytes)
    pdf_bytes = _replace_pdf_images_hq(pdf_bytes, orig_images)

    return send_file(
        io.BytesIO(pdf_bytes),
        mimetype='application/pdf',
        as_attachment=True,
        download_name='DBot_Site_Assessment.pdf'
    )


if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5000)
